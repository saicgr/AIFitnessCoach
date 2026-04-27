"""
GymEquipmentExtractor — AI Gym-Equipment Importer pipeline.

Normalizes inputs from any of these sources into a uniform equipment list:
  - source='file'   + s3_key + mime_type   (PDF / DOCX / image)
  - source='images' + s3_keys[]            (multiple images, single Gemini call)
  - source='text'   + raw_text             (paste / chat copy)
  - source='url'    + url                  (generic HTML page scrape)

Pipeline:
  1. _load_source() → bytes / text / image parts usable by Gemini.
  2. _gemini_extract() → list of raw equipment dicts
     [{raw_name, quantity, weight_range, confidence}, ...].
  3. _canonicalize() → (matched, unmatched) using EquipmentResolver.
  4. _infer_environment() → 'commercial_gym' | 'home_gym' | 'home' | 'outdoor' | 'hotel'.

Per project conventions:
  - NO mock data, NO silent fallbacks — errors surface.
  - Weight strings preserved verbatim (user works out in lbs).
  - Extensive 🏋️ / ✅ / ❌ / ⚠️ prefixed logs.
"""

from __future__ import annotations

import asyncio
import io
import json
import re
from typing import Optional

from google.genai import types

from core import branding
from core.logger import get_logger
from services.equipment_resolver import EquipmentResolver
from services.gemini.constants import gemini_generate_with_retry
from services.vision_service import VisionService, get_vision_service

logger = get_logger(__name__)


# Environment inference tiers (canonical names from equipment_types table).
# Order: if any commercial-only item appears → commercial_gym wins.
_COMMERCIAL_ONLY = {
    "smith_machine", "cable_machine", "leg_press",
    "hack_squat", "leg_extension", "leg_curl",
    "lat_pulldown", "seated_row_machine", "chest_press_machine",
    "pec_deck", "ab_crunch_machine", "assisted_pullup_machine",
}
_HOME_GYM_CORE = {
    "dumbbells", "barbell", "bench", "kettlebell", "resistance_bands",
    "squat_rack", "power_rack", "ez_curl_bar",
}
_HOME_MINIMAL = {
    "bodyweight", "pull_up_bar", "resistance_bands",
}


# ---------------- URL loader config ----------------
_URL_MAX_BYTES = 2 * 1024 * 1024          # 2 MB
_URL_TIMEOUT_SECONDS = 10
_URL_MAX_TEXT_CHARS = 100_000             # cap stripped-text size sent to Gemini

# Reject obviously unsupported URLs (auth-gated, JS-heavy, video platforms).
_REJECTED_URL_HOSTS = (
    "youtube.com", "youtu.be", "m.youtube.com",
    "instagram.com", "tiktok.com",
    "facebook.com", "twitter.com", "x.com",
    "linkedin.com", "login.",
)


class GymEquipmentExtractor:
    """Central extractor orchestrating Gemini + EquipmentResolver."""

    def __init__(
        self,
        vision_service: Optional[VisionService] = None,
        equipment_resolver: Optional[EquipmentResolver] = None,
    ):
        # The vision service owns the Gemini client + S3 download helper.
        self.vision = vision_service or get_vision_service()
        # Resolver is optional at __init__ — loaded lazily in extract() so we can
        # await the async get_instance() without fighting __init__ semantics.
        self._resolver: Optional[EquipmentResolver] = equipment_resolver

    # --------------------------------------------------------------
    # Public entry point
    # --------------------------------------------------------------
    async def extract(self, source: str, **kwargs) -> dict:
        """Run the full extraction pipeline.

        Supported kwargs per source:
          - file:   s3_key (str), mime_type (str)
          - images: s3_keys (list[str])
          - text:   raw_text (str)
          - url:    url (str)

        Returns:
          {
            "matched":    [{canonical, raw, confidence, quantity, weight_range}, ...],
            "unmatched":  [{raw, confidence}, ...],
            "inferred_environment": str,
            "total_extracted": int,
          }
        """
        if self._resolver is None:
            self._resolver = await EquipmentResolver.get_instance()

        logger.info(f"🏋️ [GymEquipmentExtractor] Starting extraction (source={source})")

        raw_items = await self._load_and_extract(source, **kwargs)
        logger.info(f"🏋️ [GymEquipmentExtractor] Gemini returned {len(raw_items)} raw items")

        matched, unmatched = self._canonicalize(raw_items)
        environment = self._infer_environment([m["canonical"] for m in matched])

        result = {
            "matched": matched,
            "unmatched": unmatched,
            "inferred_environment": environment,
            "total_extracted": len(raw_items),
        }
        logger.info(
            f"✅ [GymEquipmentExtractor] Done: {len(matched)} matched, "
            f"{len(unmatched)} unmatched, env={environment}"
        )
        return result

    # --------------------------------------------------------------
    # Step 1+2: Source loading + Gemini extraction (combined so image/PDF
    # sources can use native Gemini Parts without an intermediate OCR step).
    # --------------------------------------------------------------
    async def _load_and_extract(self, source: str, **kwargs) -> list[dict]:
        if source == "text":
            raw_text = (kwargs.get("raw_text") or "").strip()
            if not raw_text:
                raise ValueError("source='text' requires non-empty raw_text")
            return await self._gemini_extract_from_text(raw_text)

        if source == "url":
            url = (kwargs.get("url") or "").strip()
            if not url:
                raise ValueError("source='url' requires a url")
            page_text = await self._fetch_url_text(url)
            return await self._gemini_extract_from_text(page_text, context_hint=f"URL: {url}")

        if source == "file":
            s3_key = kwargs.get("s3_key")
            mime_type = kwargs.get("mime_type")
            if not s3_key or not mime_type:
                raise ValueError("source='file' requires s3_key and mime_type")
            return await self._extract_from_file(s3_key, mime_type)

        if source == "images":
            s3_keys = kwargs.get("s3_keys") or []
            if not s3_keys:
                raise ValueError("source='images' requires at least one s3_key")
            return await self._extract_from_images(s3_keys)

        raise ValueError(f"Unknown source '{source}' (expected: file|images|text|url)")

    # --- Text / URL source → single Gemini text call ---
    async def _gemini_extract_from_text(
        self, text: str, context_hint: Optional[str] = None
    ) -> list[dict]:
        if not text.strip():
            logger.warning("⚠️ [GymEquipmentExtractor] Empty text, returning []")
            return []

        # Cap text size to avoid blowing token budget on scraped HTML.
        if len(text) > _URL_MAX_TEXT_CHARS:
            logger.warning(
                f"⚠️ [GymEquipmentExtractor] Trimming text from {len(text)} to "
                f"{_URL_MAX_TEXT_CHARS} chars"
            )
            text = text[:_URL_MAX_TEXT_CHARS]

        prompt = VisionService.GYM_EQUIPMENT_EXTRACTION_PROMPT
        if context_hint:
            prompt = f"{prompt}\n\nContext: {context_hint}"
        prompt = f"{prompt}\n\nContent:\n{text}"

        try:
            response = await gemini_generate_with_retry(
                model=self.vision.model,
                contents=[prompt],
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    max_output_tokens=8000,
                    response_mime_type="application/json",
                ),
                method_name="gym_equipment_extract_text",
            )
            return _parse_gemini_equipment_json(response.text or "")
        except Exception as e:
            logger.error(f"❌ [GymEquipmentExtractor] Gemini text extraction failed: {e}", exc_info=True)
            raise

    # --- File source (PDF / DOCX / single image) ---
    async def _extract_from_file(self, s3_key: str, mime_type: str) -> list[dict]:
        logger.info(f"🏋️ [GymEquipmentExtractor] Loading file from S3: {s3_key} ({mime_type})")

        # DOCX → parse locally to plain text, then run the text pipeline
        # (Gemini has no native DOCX support).
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file_bytes = await self.vision._download_image_from_s3(s3_key)
            text = _extract_docx_text(file_bytes)
            return await self._gemini_extract_from_text(text, context_hint="source=docx")

        # PDF + image: use native Gemini multimodal via VisionService helper.
        if mime_type == "application/pdf" or mime_type.startswith("image/"):
            file_bytes = await self.vision._download_image_from_s3(s3_key)
            return await self.vision.extract_equipment_from_document(file_bytes, mime_type)

        raise ValueError(
            f"Unsupported mime_type for source='file': {mime_type}. "
            f"Allowed: application/pdf, docx, image/*"
        )

    # --- Images source (multiple photos → single combined Gemini call) ---
    async def _extract_from_images(self, s3_keys: list[str]) -> list[dict]:
        logger.info(f"🏋️ [GymEquipmentExtractor] Downloading {len(s3_keys)} images from S3")
        # Parallel S3 downloads (mirrors vision_service.analyze_food_from_s3_keys pattern).
        download_tasks = [self.vision._download_image_from_s3(key) for key in s3_keys]
        image_bytes_list = await asyncio.gather(*download_tasks)

        image_parts = [
            types.Part.from_bytes(data=data, mime_type="image/jpeg")
            for data in image_bytes_list
        ]

        try:
            response = await gemini_generate_with_retry(
                model=self.vision.model,
                contents=[VisionService.GYM_EQUIPMENT_EXTRACTION_PROMPT] + image_parts,
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    max_output_tokens=8000,
                    response_mime_type="application/json",
                ),
                method_name="gym_equipment_extract_images",
            )
            return _parse_gemini_equipment_json(response.text or "")
        except Exception as e:
            logger.error(f"❌ [GymEquipmentExtractor] Image extraction failed: {e}", exc_info=True)
            raise

    # --- URL fetch → stripped text ---
    async def _fetch_url_text(self, url: str) -> str:
        import httpx

        lowered = url.lower()
        if not (lowered.startswith("http://") or lowered.startswith("https://")):
            raise ValueError(f"Only http(s) URLs supported: {url}")
        for bad in _REJECTED_URL_HOSTS:
            if bad in lowered:
                raise ValueError(
                    f"URL host not supported ({bad}). "
                    f"Video sites, social media, and auth-gated pages are out of scope."
                )

        logger.info(f"🏋️ [GymEquipmentExtractor] Fetching URL: {url}")
        try:
            async with httpx.AsyncClient(timeout=_URL_TIMEOUT_SECONDS, follow_redirects=True) as client:
                resp = await client.get(url, headers={"User-Agent": f"{branding.APP_NAME}Bot/1.0 (+gym-equipment-importer)"})
                resp.raise_for_status()
        except Exception as e:
            logger.error(f"❌ [GymEquipmentExtractor] URL fetch failed: {e}", exc_info=True)
            raise ValueError(f"Could not fetch URL: {e}") from e

        content = resp.content[:_URL_MAX_BYTES]
        ctype = (resp.headers.get("content-type") or "").lower()
        if "html" not in ctype and "text" not in ctype:
            raise ValueError(
                f"URL content-type '{ctype}' not supported — only HTML/plain-text URLs."
            )

        try:
            html = content.decode(resp.encoding or "utf-8", errors="replace")
        except Exception:
            html = content.decode("utf-8", errors="replace")

        return _strip_html(html)

    # --------------------------------------------------------------
    # Step 3: Canonicalization via EquipmentResolver
    # --------------------------------------------------------------
    def _canonicalize(self, raw_items: list[dict]) -> tuple[list[dict], list[dict]]:
        """Split into matched/unmatched buckets.

        Matched: dedupe by canonical, keep highest confidence.
        Unmatched: keep raw + confidence (UI shows "+ Add as custom" / "skip").
        """
        assert self._resolver is not None, "Resolver must be loaded before canonicalization"

        matched_by_canonical: dict[str, dict] = {}
        unmatched: list[dict] = []

        for item in raw_items:
            raw_name = (item.get("raw_name") or "").strip()
            if not raw_name:
                continue
            confidence = float(item.get("confidence") or 0.5)
            quantity = item.get("quantity")
            weight_range = item.get("weight_range")

            canonical = self._resolver.resolve(raw_name)
            if canonical:
                existing = matched_by_canonical.get(canonical)
                if existing is None or confidence > existing["confidence"]:
                    matched_by_canonical[canonical] = {
                        "canonical": canonical,
                        "raw": raw_name,
                        "confidence": confidence,
                        "quantity": quantity,
                        "weight_range": weight_range,
                    }
            else:
                unmatched.append({
                    "raw": raw_name,
                    "confidence": confidence,
                    "quantity": quantity,
                    "weight_range": weight_range,
                })

        return list(matched_by_canonical.values()), unmatched

    # --------------------------------------------------------------
    # Step 4: Environment inference
    # --------------------------------------------------------------
    def _infer_environment(self, canonical_names: list[str]) -> str:
        """Rules:
           - Any commercial-only machine → commercial_gym
           - Only home_gym_core items    → home_gym
           - Only bodyweight / pull_up_bar → home
           - Empty or unclear            → commercial_gym (UI remains editable)
        """
        if not canonical_names:
            logger.info("🏋️ [GymEquipmentExtractor] No canonical items → defaulting env=commercial_gym")
            return "commercial_gym"

        name_set = set(canonical_names)

        if name_set & _COMMERCIAL_ONLY:
            return "commercial_gym"

        if name_set and name_set.issubset(_HOME_MINIMAL):
            return "home"

        if name_set and name_set.issubset(_HOME_GYM_CORE | _HOME_MINIMAL):
            return "home_gym"

        return "commercial_gym"


# ===================================================================
# Module-level helpers
# ===================================================================

def _parse_gemini_equipment_json(raw_text: str) -> list[dict]:
    """Strip ``` fences and parse Gemini JSON output into validated list of dicts."""
    text = (raw_text or "").strip()
    if not text:
        logger.warning("⚠️ [GymEquipmentExtractor] Gemini returned empty string")
        return []

    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)

    try:
        parsed = json.loads(text)
    except json.JSONDecodeError as e:
        logger.error(f"❌ [GymEquipmentExtractor] Invalid JSON: {text[:300]}", exc_info=True)
        raise ValueError(f"Gemini returned invalid JSON: {e}") from e

    if not isinstance(parsed, list):
        logger.warning(
            f"⚠️ [GymEquipmentExtractor] Expected list, got {type(parsed).__name__}"
        )
        return []

    cleaned: list[dict] = []
    for item in parsed:
        if not isinstance(item, dict):
            continue
        raw_name = (item.get("raw_name") or "").strip()
        if not raw_name:
            continue
        try:
            confidence = float(item.get("confidence") or 0.5)
        except (TypeError, ValueError):
            confidence = 0.5
        confidence = max(0.0, min(1.0, confidence))
        cleaned.append({
            "raw_name": raw_name,
            "quantity": item.get("quantity"),
            "weight_range": item.get("weight_range"),
            "confidence": confidence,
        })
    return cleaned


def _extract_docx_text(file_bytes: bytes) -> str:
    """Pull plain text out of a DOCX file using python-docx."""
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Add 'python-docx' to backend/requirements.txt."
        ) from e

    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"❌ [GymEquipmentExtractor] python-docx failed to open DOCX: {e}", exc_info=True)
        raise ValueError(f"Could not parse DOCX: {e}") from e

    parts: list[str] = []
    for para in document.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())
    # Also walk tables (gyms often list equipment in tables).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())

    text = "\n".join(parts)
    logger.info(f"🏋️ [GymEquipmentExtractor] DOCX → {len(text)} chars across {len(parts)} blocks")
    return text


def _strip_html(html: str) -> str:
    """Strip HTML → plain text. Prefers BeautifulSoup; falls back to regex."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        # Drop non-content tags outright.
        for tag in soup(["script", "style", "noscript", "svg", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
    except Exception as e:
        logger.warning(
            f"⚠️ [GymEquipmentExtractor] BeautifulSoup unavailable/failed ({e}); "
            f"falling back to regex strip"
        )
        # Remove script/style blocks first, then all tags.
        text = re.sub(r"<(script|style|noscript)[^>]*>.*?</\1>", " ", html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)

    # Collapse whitespace.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()
